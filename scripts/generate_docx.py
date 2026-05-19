"""Regenerate .docx papers from updated markdown sources.

Reads papers/paper{1,2,3}_*.md and generates papers/docx/Paper{1,2,3}_*.docx
with proper formatting (headings, tables, code blocks, lists).
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def md_to_docx(md_path: str, docx_path: str):
    """Convert a markdown paper to a formatted .docx file."""
    with open(md_path) as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            i += 1
            continue

        # Tables
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            # Skip separator rows (|---|---|)
            if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table, render it
            in_table = False
            if table_rows:
                n_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=n_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell in enumerate(row):
                        if c_idx < n_cols:
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell)  # strip bold
                            cell_text = re.sub(r'~~(.*?)~~', r'\1', cell_text)  # strip strikethrough
                            table.rows[r_idx].cells[c_idx].text = cell_text
                            # Bold header row
                            if r_idx == 0:
                                for run in table.rows[r_idx].cells[c_idx].paragraphs[0].runs:
                                    run.bold = True
                doc.add_paragraph()  # spacing after table
                table_rows = []

        # Headings
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Bold author line
        if line.startswith('**') and line.endswith('**') and 'Grudhanti' in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip('*').strip())
            run.bold = True
            i += 1
            continue

        # List items
        if re.match(r'^\s*[-*]\s', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            text = clean_markdown(text)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        if re.match(r'^\s*\d+\.\s', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            text = clean_markdown(text)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Regular paragraph
        text = clean_markdown(line)
        if text:
            p = doc.add_paragraph()
            # Handle bold segments
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(clean_markdown(part[2:-2]))
                    run.bold = True
                else:
                    p.add_run(clean_markdown(part))

        i += 1

    # Flush any remaining table
    if in_table and table_rows:
        n_cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=n_cols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_rows):
            for c_idx, cell in enumerate(row):
                if c_idx < n_cols:
                    table.rows[r_idx].cells[c_idx].text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell)

    doc.save(docx_path)
    print(f"  ✓ {docx_path}")


def clean_markdown(text: str) -> str:
    """Remove markdown formatting from text."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # italic
    text = re.sub(r'`(.*?)`', r'\1', text)        # inline code
    text = re.sub(r'~~(.*?)~~', r'\1', text)      # strikethrough
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # links
    return text


def main():
    print("=" * 60)
    print("  Generating .docx papers from markdown")
    print("=" * 60)
    print()

    papers = [
        ("papers/paper1_staged_retrieval.md", "papers/docx/Paper1_Staged_Retrieval.docx"),
        ("papers/paper2_semantic_routing.md", "papers/docx/Paper2_Semantic_Routing.docx"),
        ("papers/paper3_memory_systems.md", "papers/docx/Paper3_Memory_Systems.docx"),
    ]

    Path("papers/docx").mkdir(parents=True, exist_ok=True)

    for md_path, docx_path in papers:
        md_to_docx(md_path, docx_path)

    print("\n  Done! All .docx files updated.")


if __name__ == "__main__":
    main()
